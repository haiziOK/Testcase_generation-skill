"""
Document parser for design documents.
Currently supports plain text files (.txt) and Word documents (.docx).
"""
import os
from typing import Optional, List, Dict
import logging

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parse design documents to extract requirements."""

    def __init__(self):
        pass

    def parse_text_file(self, file_path: str) -> str:
        """Parse a plain text file and return its content."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"Successfully read text file: {file_path}")
            return content
        except Exception as e:
            logger.error(f"Failed to parse text file {file_path}: {e}")
            raise

    def parse_docx_file(self, file_path: str) -> str:
        """Parse a Word document (.docx) and return its content."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not installed. Install with: pip install python-docx")

        try:
            doc = docx.Document(file_path)
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content_parts.append(cell.text)

            content = "\n\n".join(content_parts)
            logger.info(f"Successfully read DOCX file: {file_path}")
            return content
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {e}")
            raise

    def parse(self, file_path: str) -> str:
        """Parse a design document based on file extension."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.txt':
            return self.parse_text_file(file_path)
        elif ext == '.pdf':
            # TODO: Implement PDF parsing
            raise NotImplementedError("PDF parsing not implemented yet")
        elif ext in ['.docx', '.doc']:
            return self.parse_docx_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def extract_requirements(self, content: str) -> List[Dict]:
        """Extract requirements from document content.

        This is a simple implementation that returns the entire content as a single requirement.
        In a real implementation, this would use NLP or pattern matching to identify specific requirements.
        """
        # For MVP, we'll treat the entire content as the requirement
        # The LLM will handle requirement extraction from this content
        return [{"id": "req_1", "text": content, "type": "functional"}]

    def parse_and_extract(self, file_path: str) -> List[Dict]:
        """Parse document and extract requirements."""
        content = self.parse(file_path)
        requirements = self.extract_requirements(content)
        return requirements